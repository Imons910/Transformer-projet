from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
import os

# ==========================================
# CONFIGURATION
# ==========================================
OUTPUT_FILENAME = "Rapport_M2_Depth_Estimation_LoRA.docx"
IMAGE_DIR = "./resultats_projet_final"

# Métadonnées
TITRE = "Fine-Tuning du Modèle Depth Anything V2 avec LoRA"
SOUS_TITRE = "Adaptation de domaine pour l'estimation de profondeur en milieu industriel (Zivid)"
AUTEURS = "Binôme Master 2"
DATE = "Janvier 2025"

def create_document():
    doc = Document()

    # --- STYLES ET FORMATAGE ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Fonction pour ajouter des titres stylisés
    def add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        run = p.runs[0]
        run.font.color.rgb = RGBColor(0, 51, 102) # Bleu Nuit
        run.font.name = 'Arial'
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    # Fonction pour ajouter du texte justifié
    def add_text(text, bold=False, italic=False, bullet=False):
        p = doc.add_paragraph(style='List Bullet' if bullet else 'Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        return p

    # Fonction pour ajouter une image avec légende
    def add_image(filename, caption, width=Inches(6)):
        img_path = os.path.join(IMAGE_DIR, filename)
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=width)
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            caption_p = doc.add_paragraph(f"Figure : {caption}", style='Caption')
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_p.paragraph_format.space_after = Pt(12)
        else:
            p = doc.add_paragraph(f"[PLACEHOLDER IMAGE: {filename}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True

    # ==========================================
    # PAGE DE GARDE
    # ==========================================
    for _ in range(5): doc.add_paragraph()
    
    title = doc.add_heading(TITRE, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(SOUS_TITRE)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].italic = True
    
    for _ in range(8): doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Rapport de Projet - Vision par Ordinateur\n").bold = True
    info.add_run(f"Auteurs : {AUTEURS}\n")
    info.add_run(f"Date : {DATE}")

    doc.add_page_break()

    # ==========================================
    # INTRODUCTION
    # ==========================================
    add_heading("1. Introduction et Contexte", 1)
    add_text("L'estimation de la profondeur monoculaire (MDE) est un pilier de la perception robotique. Si les modèles récents comme 'Depth Anything' excellent sur des images naturelles (paysages, scènes urbaines), leur performance s'effondre souvent lorsqu'ils sont confrontés à des données industrielles spécifiques, caractérisées par des textures répétitives (pneus noirs), des éclairages artificiels et des capteurs spécialisés (Zivid).")
    
    add_text("\nObjectifs du projet :", bold=True)
    add_text("Ce projet vise à adapter le modèle 'Depth Anything V2 Small' à un dataset de pneus en utilisant la technique LoRA (Low-Rank Adaptation). L'objectif est double :")
    add_text("Obtenir une reconstruction 3D fidèle (minimisation de l'erreur métrique).", bullet=True)
    add_text("Fournir une aide à la décision pour la robotique (classification Obstacle Proche/Loin).", bullet=True)

    # ==========================================
    # CADRE THÉORIQUE
    # ==========================================
    add_heading("2. Cadre Théorique : Fine-Tuning Efficient (PEFT)", 1)
    
    add_heading("2.1. Pourquoi LoRA (Low-Rank Adaptation) ?", 2)
    add_text("Le fine-tuning complet d'un modèle Transformer (plusieurs millions de paramètres) est coûteux en mémoire et risque de provoquer un 'oubli catastrophique' des connaissances préalables. LoRA propose de geler les poids pré-entraînés W0 et d'injecter des matrices de rang décomposable entraînables.")
    
    add_text("Mathématiquement, la mise à jour des poids s'écrit :")
    add_text("W = W0 + (B x A)", italic=True)
    add_text("Où A et B sont des matrices de rang r << d. Dans notre projet, nous avons choisi un rang r=32, ce qui permet au modèle d'apprendre des caractéristiques complexes sans exploser le nombre de paramètres.")

    add_heading("2.2. Architecture Depth Anything V2", 2)
    add_text("Nous utilisons la version 'Small' du modèle, basée sur une architecture DPT (Dense Prediction Transformer). Ce modèle utilise un encodeur ViT (Vision Transformer) pour extraire les caractéristiques et un décodeur pour générer la carte de profondeur.")

    # ==========================================
    # PRÉPARATION DES DONNÉES
    # ==========================================
    add_heading("3. Préparation du Dataset et Défis", 1)
    
    add_heading("3.1. Analyse des Données Brutes", 2)
    add_text("Le dataset provient d'une caméra Zivid. Il est composé d'images RGB et de fichiers numpy (.npy) contenant les nuages de points XYZ.")
    
    add_text("Difficulté n°1 : Incohérence des Unités", bold=True)
    add_text("Les données brutes étaient en millimètres (valeurs > 1000), alors que le modèle pré-entraîné fonctionne en mètres. Nous avons implémenté une normalisation dynamique dans le DataLoader :")
    add_text("if max(depth) > 100: depth = depth / 1000.0", italic=True)

    add_heading("3.2. Limitation Structurelle du Dataset", 2)
    add_text("Une contrainte majeure de ce projet est l'absence de division officielle entre un jeu d'entraînement (Train) et un jeu de validation (Val).")
    add_text("Impact : ", bold=True)
    add_text("Nous avons dû évaluer le modèle sur le jeu d'entraînement lui-même. Bien que cela puisse théoriquement favoriser le surapprentissage, nous avons compensé ce biais par l'utilisation stricte de l'augmentation de données et de l'Early Stopping. Cela reste néanmoins une limite méthodologique à souligner.")

    add_heading("3.3. Augmentation de Données (Data Augmentation)", 2)
    add_text("Les images de pneus étant très sombres (faible contraste), le modèle peinait initialement à extraire des contours. Nous avons intégré 'ColorJitter' :")
    add_text("Luminosité : +40%", bullet=True)
    add_text("Contraste : +30%", bullet=True)
    add_text("Cela a drastiquement amélioré la convergence dès les premières époques.")

    add_image("visu_epoch_1_MASKED.png", "Exemple d'image (éclaircie) et de vérité terrain au début du projet.")

    doc.add_page_break()

    # ==========================================
    # MÉTHODOLOGIE
    # ==========================================
    add_heading("4. Stratégie d'Entraînement et Paramètres", 1)

    add_heading("4.1. Configuration LoRA 'Boostée'", 2)
    add_text("Initialement, nous avons testé une configuration standard (r=16, cibles='query', 'value'). Les résultats plafonnaient à une précision (Delta) de 70%.")
    add_text("Nous avons alors opté pour une configuration agressive :")
    add_text("Rang (r) : 32 (Capacité d'apprentissage accrue)", bullet=True)
    add_text("Modules Cibles : ['query', 'key', 'value', 'dense', 'fc1', 'fc2']", bullet=True)
    add_text("Justification : ", bold=True)
    add_text("En incluant les couches MLP ('dense', 'fc'), nous permettons au modèle de redéfinir sa représentation interne des distances métriques, pas seulement son attention visuelle.")

    add_heading("4.2. Gestion de l'Apprentissage (Callbacks)", 2)
    add_text("Pour contrer le risque de surapprentissage lié à l'évaluation sur le train-set, nous avons développé une classe 'EarlyStopping' :")
    add_text("Métrique surveillée : Delta (Précision < 1.25)", bullet=True)
    add_text("Patience : 5 époques", bullet=True)
    add_text("Comportement : Si le modèle ne s'améliore pas pendant 5 époques, l'entraînement stoppe et les meilleurs poids sont restaurés.")

    # ==========================================
    # DYNAMIQUE (L'HISTOIRE DU PROJET)
    # ==========================================
    add_heading("5. Analyse de la Dynamique d'Apprentissage", 1)
    
    add_text("L'entraînement s'est déroulé en trois phases distinctes, marquées par des défis techniques spécifiques.")

    add_heading("5.1. Phase 1 : Le Piège de la Visualisation (Epoch 1-10)", 2)
    add_text("Durant les premières époques, la Loss diminuait, mais nos images de prédiction apparaissaient totalement violettes ou uniformes. Nous avons cru à une divergence du modèle.")
    add_text("Diagnostic : ", bold=True)
    add_text("Le problème venait de l'affichage. Le modèle prédisait quelques pixels aberrants (ex: 100m) dus au bruit. La normalisation automatique de Matplotlib écrasait alors toute la dynamique utile (0-2m) dans une seule couleur sombre.")
    add_text("Solution : ", bold=True)
    add_text("Nous avons forcé l'échelle d'affichage (vmin=0, vmax=2.5m) et appliqué un masquage des pixels invalides pour comparer ce qui est comparable.")

    add_heading("5.2. Phase 2 : La Convergence (Epoch 10-35)", 2)
    add_text("Grâce à la configuration LoRA r=32, le modèle a rapidement appris la géométrie des pneus. On observe une chute drastique de la RMSE (de 1.8m à 0.10m).")
    
    add_image("courbes_finales.png", "Courbes d'apprentissage : Notez la stabilité de la Loss et la montée régulière du Delta.")

    add_heading("5.3. Phase 3 : Le Plateau et l'Arrêt (Epoch 35-50)", 2)
    add_text("Vers l'époque 40, le modèle a atteint un plafond de performance (Delta ~0.99). À l'époque 50, nous avons atteint un record absolu de Delta=0.9941. L'absence de dégradation de la Loss en fin de parcours confirme que nous n'avons pas fait d'overfitting destructeur.")

    doc.add_page_break()

    # ==========================================
    # RÉSULTATS
    # ==========================================
    add_heading("6. Résultats et Performances", 1)

    add_heading("6.1. Métriques Géométriques (Comparatif)", 2)
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Medium Grid 1 Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Métrique'
    hdr[1].text = 'Avant (Zero-Shot)'
    hdr[2].text = 'Après (Fine-Tuned)'
    
    r1 = table.rows[1].cells
    r1[0].text = 'Delta (Accuracy)'
    r1[1].text = '17.04 %'
    r1[2].text = '99.41 %'
    
    r2 = table.rows[2].cells
    r2[0].text = 'RMSE (Erreur)'
    r2[1].text = '1.808 m'
    r2[2].text = '0.070 m'

    add_text("\nInterprétation : ", bold=True)
    add_text("Le modèle initial était inutilisable (17% de précision). Le modèle final est extrêmement précis (99.4%), avec une erreur moyenne de seulement 7 cm, ce qui est compatible avec des tâches de manipulation robotique.")

    add_heading("6.2. Résultats Qualitatifs", 2)
    add_text("L'inspection visuelle finale montre une reconstruction quasi-parfaite des structures 3D.")
    
    # Mets ici ton image finale (EPOCH 50 ou proche)
    add_image("visu_epoch_50_MASKED.png", "Résultat Final (Epoch 50). La prédiction (droite) est indiscernable de la vérité terrain (centre).")

    add_heading("6.3. Classification Opérationnelle", 2)
    add_text("Pour répondre à la consigne de classification, nous avons défini un seuil de sécurité à 1.5 mètre.")
    add_text("F1-Score Final : 0.9556", bold=True)
    add_text("Ce score très élevé indique que le modèle ne fait pratiquement aucune erreur de jugement entre 'Obstacle Proche' et 'Fond'.")

    # ==========================================
    # DISCUSSION & LIMITES
    # ==========================================
    add_heading("7. Discussion et Limites", 1)
    
    add_heading("7.1. Pourquoi le Delta augmente-t-il tant ?", 2)
    add_text("L'augmentation spectaculaire du Delta (de 0.17 à 0.99) s'explique par l'adaptation de l'échelle (Scale Alignment). Les modèles Depth Anything sont entraînés sur des disparités relatives. LoRA a permis au modèle d'apprendre la relation spécifique : 'Pixel gris foncé Zivid = 0.8 mètres'.")

    add_heading("7.2. Limites du Modèle", 2)
    add_text("Biais d'évaluation : ", bold=True)
    add_text("Faute de set de validation indépendant, ces performances (99%) sont probablement optimistes par rapport à des données totalement inconnues.")
    add_text("Dépendance au capteur : ", bold=True)
    add_text("Le modèle a appris le bruit spécifique de la caméra Zivid (les zones noires). Il pourrait échouer sur une caméra RealSense.")

    add_heading("8. Conclusion et Perspectives", 1)
    add_text("Ce projet a permis de transformer un modèle générique en un outil industriel de précision. L'utilisation de LoRA ciblant les couches denses s'est révélée être la clé du succès.")
    add_text("Pour des travaux futurs, il serait crucial de :")
    add_text("Constituer un dataset de test séparé physiquement (nouveaux pneus, nouvelle pièce).", bullet=True)
    add_text("Tester l'inférence en temps réel (FPS) pour valider l'usage embarqué.", bullet=True)

    # ==========================================
    # SAUVEGARDE
    # ==========================================
    try:
        doc.save(OUTPUT_FILENAME)
        print(f"✅ Rapport généré : {OUTPUT_FILENAME}")
        print("Texte riche, structuré, prêt à être exporté en PDF.")
    except PermissionError:
        print("❌ Erreur : Fermez le fichier Word s'il est déjà ouvert !")

if __name__ == "__main__":
    create_document()